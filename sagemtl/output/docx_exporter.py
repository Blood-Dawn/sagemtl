"""DOCX exporter for SageMTL."""

from __future__ import annotations

from typing import List

from .base import (
    BaseExporter,
    ExportCapability,
    ExportConfig,
    ExportResult,
    NovelData,
)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocxExporter(BaseExporter):
    """Export novel to DOCX (Microsoft Word) format.

    Features:
    - Styled document with headings
    - Cover page
    - Table of contents
    - Custom fonts and sizes
    """

    @classmethod
    def name(cls) -> str:
        return "Microsoft Word"

    @classmethod
    def extensions(cls) -> List[str]:
        return ["docx"]

    @classmethod
    def capabilities(cls) -> ExportCapability:
        return (
            ExportCapability.COVER |
            ExportCapability.TOC |
            ExportCapability.METADATA |
            ExportCapability.STYLES |
            ExportCapability.VOLUMES
        )

    @classmethod
    def is_available(cls) -> bool:
        return HAS_DOCX

    def export(
        self,
        novel: NovelData,
        config: ExportConfig,
    ) -> ExportResult:
        """Export novel to DOCX file."""
        if not HAS_DOCX:
            return ExportResult.failure(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )

        output_path = self._get_output_path(novel, config)

        try:
            doc = Document()

            # Set up styles
            self._setup_styles(doc, config)

            # Cover page
            if config.include_cover:
                self._add_cover_page(doc, novel, config)

            # Table of contents
            if config.include_toc:
                self._add_toc(doc, novel)

            # Chapters
            current_volume = None
            for chapter in novel.chapters:
                # Volume header
                if chapter.volume_index is not None and chapter.volume_index != current_volume:
                    current_volume = chapter.volume_index
                    volume = next(
                        (v for v in novel.volumes if v.index == current_volume),
                        None
                    )
                    if volume:
                        if config.volume_page_break:
                            doc.add_page_break()
                        heading = doc.add_heading(volume.title, level=1)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Chapter
                doc.add_heading(chapter.title, level=2)

                # Add content
                for paragraph in chapter.content.split('\n\n'):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.style = 'Body Text'

                # Page break between chapters (optional)
                if config.volume_page_break:
                    doc.add_page_break()

            # Set document properties
            core_props = doc.core_properties
            core_props.title = config.title or novel.title
            if novel.author:
                core_props.author = config.author or novel.author
            if novel.description:
                core_props.comments = novel.description

            # Save document
            doc.save(str(output_path))

            return ExportResult.ok(output_path)

        except Exception as e:
            return ExportResult.failure(str(e))

    def _setup_styles(self, doc: Document, config: ExportConfig) -> None:
        """Set up document styles."""
        styles = doc.styles

        # Title style
        try:
            title_style = styles['Title']
        except KeyError:
            title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)

        title_style.font.name = config.font_family
        title_style.font.size = Pt(28)
        title_style.font.bold = True

        # Heading 1 (Volume)
        try:
            h1_style = styles['Heading 1']
        except KeyError:
            h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

        h1_style.font.name = config.font_family
        h1_style.font.size = Pt(20)
        h1_style.font.bold = True

        # Heading 2 (Chapter)
        try:
            h2_style = styles['Heading 2']
        except KeyError:
            h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)

        h2_style.font.name = config.font_family
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True

        # Body text
        try:
            body_style = styles['Body Text']
        except KeyError:
            body_style = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)

        body_style.font.name = config.font_family
        body_style.font.size = Pt(config.font_size)
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        body_style.paragraph_format.space_after = Pt(6)

    def _add_cover_page(
        self,
        doc: Document,
        novel: NovelData,
        config: ExportConfig,
    ) -> None:
        """Add cover page."""
        title = config.title or novel.title
        author = config.author or novel.author

        # Add cover image if available
        if novel.cover_path and novel.cover_path.exists():
            try:
                doc.add_picture(str(novel.cover_path), width=Inches(4))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass  # Skip if image fails

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(32)
        run.font.bold = True

        # Author
        if author:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"by {author}")
            run.font.size = Pt(18)
            run.font.italic = True

        # Description
        if novel.description:
            doc.add_paragraph()
            p = doc.add_paragraph(novel.description)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _add_toc(self, doc: Document, novel: NovelData) -> None:
        """Add table of contents."""
        heading = doc.add_heading("Table of Contents", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_volume = None
        for chapter in novel.chapters:
            # Volume header
            if chapter.volume_index is not None and chapter.volume_index != current_volume:
                current_volume = chapter.volume_index
                volume = next(
                    (v for v in novel.volumes if v.index == current_volume),
                    None
                )
                if volume:
                    p = doc.add_paragraph()
                    run = p.add_run(volume.title)
                    run.font.bold = True

            # Chapter entry
            p = doc.add_paragraph(f"    {chapter.title}")

        doc.add_page_break()
